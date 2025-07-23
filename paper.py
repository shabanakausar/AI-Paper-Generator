import os
from flask import Flask, render_template_string, request, redirect, url_for, session, send_file
from werkzeug.utils import secure_filename
from langchain_community.document_loaders import PyPDFLoader, TextLoader, Docx2txtLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.chains import RetrievalQA
from langchain_groq import ChatGroq
from dotenv import load_dotenv
from docx import Document
from fpdf import FPDF

# Load API key
load_dotenv()

# ---- Flask setup ----
app = Flask(__name__)
app.secret_key = 'super-secret-key'  # for session handling
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['GENERATED_FOLDER'] = 'generated'
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['GENERATED_FOLDER'], exist_ok=True)

# ---- Global variables ----
vectorstore = None

# ---- Dummy user login ----
USERS = {
    "teacher": "pass123"
}

# ---- Templates ----

LOGIN_TEMPLATE = """
<!DOCTYPE html>
<html>
<head><title>Login</title>
<link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.3/dist/css/bootstrap.min.css" rel="stylesheet"></head>
<body class="bg-light">
<div class="container py-5">
    <h2>🔐 Login</h2>
    {% if error %}<div class="alert alert-danger">{{ error }}</div>{% endif %}
    <form method="POST">
        <input class="form-control mb-2" type="text" name="username" placeholder="Username" required>
        <input class="form-control mb-2" type="password" name="password" placeholder="Password" required>
        <button class="btn btn-primary">Login</button>
    </form>
</div></body></html>
"""

HOME_TEMPLATE = """
<!DOCTYPE html>
<html>
<head>
    <title>AI Teacher Assistant</title>
    <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.3/dist/css/bootstrap.min.css" rel="stylesheet">
</head>
<body class="bg-light">
<div class="container py-5">
    <div class="d-flex justify-content-between">
        <h2>📘 AI Teacher Assistant</h2>
        <a href="/logout" class="btn btn-outline-danger">Logout</a>
    </div>

    {% if message %}
        <div class="alert alert-info">{{ message }}</div>
    {% endif %}

    <form method="POST" enctype="multipart/form-data" class="mb-4">
        <div class="mb-3">
            <label class="form-label">Upload Document(s)</label>
            <input type="file" name="files" multiple class="form-control" required>
        </div>
        <button type="submit" class="btn btn-primary">Upload & Process</button>
    </form>

    <div class="row g-4">
        {% for qtype in ["short", "long", "mcq"] %}
        <div class="col-md-4">
            <h5>Generate {{ qtype.capitalize() }} Questions</h5>
            <form action="/generate" method="POST">
                <input type="hidden" name="type" value="{{ qtype }}">
                <div class="mb-2">
                    <label>Number of Questions</label>
                    <input type="number" name="count" class="form-control" min="1" max="10" value="5">
                </div>
                <div class="mb-2">
                    <label>File Format</label>
                    <select name="filetype" class="form-select">
                        <option value="docx">DOCX</option>
                        <option value="pdf">PDF</option>
                    </select>
                </div>
                <button type="submit" class="btn btn-success">Generate</button>
            </form>
        </div>
        {% endfor %}
    </div>
</div>
</body>
</html>
"""

RESULT_TEMPLATE = """
<!DOCTYPE html>
<html>
<head>
    <title>Generated Questions</title>
    <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.3/dist/css/bootstrap.min.css" rel="stylesheet">
</head>
<body class="bg-light p-5">
<div class="container">
    <h2 class="mb-4">✅ Generated {{ q_type }} Questions</h2>
    <pre class="bg-white p-3 border rounded">{{ result }}</pre>
    <a href="{{ download_link }}" class="btn btn-success mt-3">⬇ Download as {{ filetype.upper() }}</a>
    <a href="/" class="btn btn-secondary mt-3">← Back</a>
</div>
</body>
</html>
"""

# ---- Helper Functions ----

def load_and_split(file_path, ext):
    if ext == "pdf":
        loader = PyPDFLoader(file_path)
    elif ext == "txt":
        loader = TextLoader(file_path)
    elif ext == "docx":
        loader = Docx2txtLoader(file_path)
    else:
        return []
    docs = loader.load()
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=150)
    return splitter.split_documents(docs)

def build_vectorstore(docs):
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    return FAISS.from_documents(docs, embeddings)

def generate_questions(q_type, num):
    llm = ChatGroq(
        model_name="llama3-8b-8192",
        temperature=0.3,
        api_key=os.getenv("GROQ_API_KEY")
    )
    qa = RetrievalQA.from_chain_type(llm=llm, retriever=vectorstore.as_retriever())

    if q_type == "short":
        prompt = f"Generate {num} short answer questions for a test."
    elif q_type == "long":
        prompt = f"Generate {num} long answer questions for a test."
    elif q_type == "mcq":
        prompt = f"Generate {num} multiple choice questions with 4 options each. Clearly mark the correct answer."
    else:
        prompt = "Generate questions."

    return qa.run(prompt)

def save_to_docx(content, filename):
    doc = Document()
    doc.add_heading('Generated Questions', 0)
    doc.add_paragraph(content)
    path = os.path.join(app.config['GENERATED_FOLDER'], filename)
    doc.save(path)
    return path

def save_to_pdf(content, filename):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    for line in content.split('\n'):
        pdf.multi_cell(0, 10, line)
    path = os.path.join(app.config['GENERATED_FOLDER'], filename)
    pdf.output(path)
    return path

# ---- Routes ----

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        u = request.form.get("username")
        p = request.form.get("password")
        if USERS.get(u) == p:
            session['user'] = u
            return redirect('/')
        else:
            return render_template_string(LOGIN_TEMPLATE, error="Invalid credentials")
    return render_template_string(LOGIN_TEMPLATE)

@app.route('/logout')
def logout():
    session.clear()
    return redirect('/login')

@app.route('/', methods=['GET', 'POST'])
def index():
    global vectorstore
    if 'user' not in session:
        return redirect('/login')

    message = None
    if request.method == 'POST':
        files = request.files.getlist('files')
        all_docs = []
        for file in files:
            filename = secure_filename(file.filename)
            ext = filename.split('.')[-1].lower()
            save_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(save_path)
            all_docs.extend(load_and_split(save_path, ext))
        if all_docs:
            vectorstore = build_vectorstore(all_docs)
            message = f"✅ Processed {len(files)} file(s) successfully."
        else:
            message = "⚠️ No valid content found."
    return render_template_string(HOME_TEMPLATE, message=message)

@app.route('/generate', methods=['POST'])
def generate():
    global vectorstore
    if 'user' not in session:
        return redirect('/login')
    if not vectorstore:
        return "<h3 style='color:red;'>⚠️ Upload and process a document first.</h3>"

    q_type = request.form.get("type")
    count = int(request.form.get("count"))
    filetype = request.form.get("filetype", "pdf")

    result = generate_questions(q_type, count)
    filename = f"{q_type}_questions.{filetype}"

    if filetype == "docx":
        filepath = save_to_docx(result, filename)
    else:
        filepath = save_to_pdf(result, filename)

    return render_template_string(
        RESULT_TEMPLATE,
        result=result,
        q_type=q_type.capitalize(),
        filetype=filetype,
        download_link=f"/download/{filename}"
    )

@app.route('/download/<filename>')
def download(filename):
    return send_file(os.path.join(app.config['GENERATED_FOLDER'], filename), as_attachment=True)

# ---- Start Server ----

if __name__ == '__main__':
    app.run(debug=True, use_reloader=False)
